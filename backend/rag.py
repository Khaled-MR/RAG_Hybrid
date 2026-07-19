"""
End-to-end RAG pipeline.

Flow (query):
    user query
      → embed
      → hybrid search (vector + BM25 with RRF) → top-20 candidates
      → cross-encoder rerank → top-5
      → LLM generates answer grounded in the top-5

Flow (ingest):
    text → recursive chunking → batch embed → store in LanceDB
    (call build_indexes() once after all documents are ingested)
"""

import json
import re
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

# Silence the harmless "XLMRobertaTokenizerFast ... use __call__" notices that
# transformers prints when FlagEmbedding loads the embedder/reranker tokenizers.
from transformers.utils import logging as _hf_logging
_hf_logging.set_verbosity_error()

from config import RAGConfig
from chunking import RecursiveChunker
from embeddings import BGEEmbedder
from reranker import BGEReranker
from qdrant_store import QdrantStore
from llm import OllamaLLM


class RAGPipeline:
    def __init__(self, config: Optional[RAGConfig] = None):
        self.config = config or RAGConfig()

        # If the config asks for CUDA but there's no working CUDA build of
        # torch, fall back to CPU instead of crashing on model load.
        embed_device = self._resolve_device(self.config.embedding_device)
        rerank_device = self._resolve_device(self.config.reranker_device)
        embed_fp16 = self.config.embedding_use_fp16 and embed_device != "cpu"
        rerank_fp16 = self.config.reranker_use_fp16 and rerank_device != "cpu"

        self.chunker = RecursiveChunker(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
        )
        self.embedder = BGEEmbedder(
            model_name=self.config.embedding_model,
            use_fp16=embed_fp16,
            device=embed_device,
        )
        self.reranker = BGEReranker(
            model_name=self.config.reranker_model,
            use_fp16=rerank_fp16,
            device=rerank_device,
        )
        self.store = QdrantStore(
            db_path=self.config.qdrant_path,
            collection_name=self.config.collection_name,
            embedding_dim=self.config.embedding_dim,
            url=self.config.qdrant_url,
        )
        self.store.create_or_open()
        self.llm = OllamaLLM(
            model=self.config.llm_model,
            base_url=self.config.ollama_base_url,
            keep_alive=self.config.llm_keep_alive,
            num_ctx=self.config.num_ctx,
        )

    

    def ingest_text(
        self,
        text: str,
        source: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> int:
        chunks = self.chunker.split_text(text)
        if not chunks:
            return 0

        dense, sparse = self.embedder.embed_documents_hybrid(
            chunks,
            batch_size=self.config.embedding_batch_size,
        )

        meta_json = json.dumps(metadata or {}, ensure_ascii=False)
        records = [
            {
                "id": str(uuid.uuid4()),
                "text": chunk,
                "source": source,
                "metadata": meta_json,
                "dense": vec.tolist(),
                "sparse": sp,
                "seq": i,                            # position within this file
                "section": self._detect_section(chunk),  # e.g. "المادة 17"
                "article_nums": self._article_numbers(chunk),  # every article # in chunk
            }
            for i, (chunk, vec, sp) in enumerate(zip(chunks, dense, sparse))
        ]
        self.store.add(records)
        return len(records)

    # Below this much total VRAM, "auto" keeps the embedder/reranker on the CPU
    # so the whole GPU is free for the LLM (which is the real bottleneck).
    # Otherwise a 7-9B model + the two BGE models won't all fit and the LLM
    # spills to CPU -> generation drops from ~80 tok/s to ~8 tok/s.
    _AUTO_GPU_VRAM_GB = 10

    @staticmethod
    def _resolve_device(device: str) -> str:
        import sys

        try:
            import torch

            cuda_ok = torch.cuda.is_available()
        except ImportError:
            cuda_ok = False

        if device == "auto":
            if not cuda_ok:
                return "cpu"
            vram_gb = torch.cuda.get_device_properties(0).total_memory / 1e9
            if vram_gb < RAGPipeline._AUTO_GPU_VRAM_GB:
                print(
                    f"[info] {vram_gb:.0f}GB GPU detected — keeping embedder/"
                    f"reranker on CPU so the LLM gets the full GPU.",
                    file=sys.stderr,
                )
                return "cpu"
            return "cuda"

        if device.startswith("cuda") and not cuda_ok:
            print(
                "[warn] device 'cuda' requested but CUDA is not available; "
                "falling back to CPU. Install a CUDA build of torch for GPU "
                "acceleration.",
                file=sys.stderr,
            )
            return "cpu"
        return device

    # Extensions we know how to read.
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".md", ".csv"}

    def ingest_file(self, file_path: str) -> int:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text = self._read_pdf(path)
        elif suffix == ".docx":
            text = self._read_docx(path)
        elif suffix in (".xlsx", ".xls"):
            text = self._read_excel(path)
        elif suffix == ".csv":
            text = self._read_csv(path)
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
        return self.ingest_text(
            text,
            source=str(path),
            metadata={"filename": path.name},
        )

    @staticmethod
    def _read_docx(path: Path) -> str:
        """Extract paragraphs and table cells from a .docx (Word) file."""
        from docx import Document

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    @staticmethod
    def _read_pdf(path: Path) -> str:
        # PyMuPDF (fitz) is ~5-10x faster than pypdf on large PDFs. Fall back
        # to pypdf if it isn't installed.
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(str(path))
            pages = [page.get_text() for page in doc]
            doc.close()
            return "\n\n".join(pages)
        except ImportError:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = [(page.extract_text() or "") for page in reader.pages]
            return "\n\n".join(pages)

    @staticmethod
    def _read_excel(path: Path) -> str:
        """Flatten every sheet to text: 'col: value | col: value' per row."""
        import pandas as pd

        sheets = pd.read_excel(path, sheet_name=None, dtype=str, engine=None)
        parts: List[str] = []
        for sheet_name, df in sheets.items():
            df = df.fillna("")
            parts.append(f"### Sheet: {sheet_name}")
            for _, row in df.iterrows():
                cells = [f"{col}: {val}" for col, val in row.items() if str(val).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    def _read_csv(path: Path) -> str:
        import pandas as pd

        df = pd.read_csv(path, dtype=str).fillna("")
        rows = [
            " | ".join(f"{col}: {val}" for col, val in row.items() if str(val).strip())
            for _, row in df.iterrows()
        ]
        return "\n".join(rows)

    def build_indexes(self) -> None:
        """
        Call once after ingesting all documents. Qdrant maintains its dense and
        sparse indexes automatically, so this is a no-op kept for API parity.
        """
        self.store.build_fts_index("text")
        self.store.build_vector_index()

    # ---------- Query rewrite & citation formatting (quality) ----------

    def rewrite_query(self, question, history=None):
        """Colloquial/follow-up question -> formal STANDALONE query for retrieval."""
        if not self.config.enable_query_rewrite:
            return question
        return self.llm.rewrite(question, history=history)

    # Detect an article/section label so citations & filtering can be exact.
    _ARTICLE_RE = re.compile(
        r"(?:المادة|مادة|الفصل|الباب|القسم|البند)\s+[\d٠-٩]+"
        r"|(?:Article|Section|Clause|Chapter)\s+\d+",
        re.IGNORECASE,
    )

    @staticmethod
    def _normalize_section(s: str) -> str:
        # unify Arabic-Indic digits and whitespace so query & stored labels match
        trans = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
        return " ".join(s.translate(trans).split()).lower()

    def _detect_section(self, text: str) -> str:
        m = self._ARTICLE_RE.search((text or "")[:120])
        return self._normalize_section(m.group(0)) if m else ""

    def _query_section(self, query: str) -> str:
        """If the user's query names an article ('المادة 17'), return it normalized."""
        m = self._ARTICLE_RE.search(query or "")
        return self._normalize_section(m.group(0)) if m else ""

    # Pull EVERY article number mentioned in a text, normalized to int.
    _ARTICLE_NUM_RE = re.compile(
        r"(?:المادة|مادة|المـادة|Article|Art\.?)\s*\(?\s*([\d٠-٩]{1,3})",
        re.IGNORECASE,
    )

    @classmethod
    def _article_numbers(cls, text: str):
        trans = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
        nums = set()
        for m in cls._ARTICLE_NUM_RE.finditer(text or ""):
            try:
                nums.add(int(m.group(1).translate(trans)))
            except ValueError:
                pass
        return sorted(nums)

    @classmethod
    def _query_article_num(cls, query: str):
        """The article number the user is asking about, or None."""
        nums = cls._article_numbers(query)
        return nums[0] if nums else None

    def format_contexts(self, retrieved: List[Dict[str, Any]]) -> List[str]:
        """
        Label each passage with its source file and detected article so the LLM
        can cite precisely as [المصدر: <file> - المادة N].
        """
        out = []
        for r in retrieved:
            source = Path(r.get("source", "unknown")).name
            m = self._ARTICLE_RE.search((r.get("text", "") or "")[:120])
            label = f"المصدر: {source}" + (f" - {m.group(0)}" if m else "")
            out.append(f"({label})\n{r.get('text', '')}")
        return out

    def validate_citations(self, answer: str, retrieved: List[Dict[str, Any]]):
        """Anti-hallucination gate (policy 8.5): every article number cited in the
        answer MUST appear in the retrieved context. Returns the list of article
        numbers cited but NOT supported by any retrieved chunk (= hallucinated)."""
        cited = set(self._article_numbers(answer))
        if not cited:
            return []
        supported = set()
        for r in retrieved:
            supported.update(self._article_numbers(r.get("text", "")))
        return sorted(cited - supported)

    # ---------- Retrieval ----------

    def _hybrid_one(self, q: str) -> List[Dict[str, Any]]:
        qd, qs = self.embedder.embed_query_hybrid(q)
        return self.store.hybrid_search(
            query_dense=qd, query_sparse=qs,
            top_k=self.config.initial_top_k,
            vector_weight=self.config.vector_weight,
            bm25_weight=self.config.bm25_weight,
            rrf_k=self.config.rrf_k,
        )

    def retrieve(self, query: str, extra_queries: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        # Stage 1: hybrid search for each query variant, RRF-merge (multi-query).
        queries = [query] + [q for q in (extra_queries or []) if q and q != query]
        merged: Dict[Any, Dict] = {}
        for q in queries:
            for rank, item in enumerate(self._hybrid_one(q), start=1):
                m = merged.get(item["id"])
                if m is None:
                    merged[item["id"]] = {"item": item, "score": 0.0}
                    m = merged[item["id"]]
                m["score"] += 1.0 / (self.config.rrf_k + rank)
        candidates = [m["item"] for m in
                      sorted(merged.values(), key=lambda x: x["score"], reverse=True)]
        candidates = candidates[: self.config.initial_top_k]

        # Article boost: if the query names an article (e.g. "المادة 38"), pull
        # the chunks that ACTUALLY mention that number to the very top — this is
        # what stops "ask about article 38, get article 49" retrieval errors.
        exact_ids: set = set()
        if self.config.enable_article_filter:
            num = self._query_article_num(query)
            if num is not None:
                exact = self.store.search_by_article(num, top_k=8)
                if exact:
                    exact_ids = {c["id"] for c in exact}
                    candidates = exact + [c for c in candidates if c["id"] not in exact_ids]

        if not candidates:
            return []

        # Stage 2: rerank → top final_top_k
        reranked = self.reranker.rerank(
            query=query,
            documents=[c["text"] for c in candidates],
            top_k=self.config.final_top_k,
        )
        result = [{**candidates[idx], "rerank_score": float(score)} for idx, score in reranked]

        # If the user asked for a specific article, GUARANTEE its exact chunks are
        # in the final context (the reranker must not drop the requested article).
        if exact_ids:
            present = {r["id"] for r in result}
            forced = [c for c in candidates if c["id"] in exact_ids and c["id"] not in present]
            if forced:
                result = forced[:2] + result           # prepend the real article
                result = result[: self.config.final_top_k]
        return result

    def _expand_with_neighbors(self, retrieved: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Add neighbouring chunks (same file, seq +/- window) for fuller context.
        Returns an ordered, de-duplicated list (primary hits + neighbours)."""
        if not self.config.enable_neighbor_expansion:
            return retrieved
        # Always keep the primary (reranked) hits; add neighbours up to the cap
        # so the prompt stays small enough to keep generation fast.
        by_id = {r["id"]: r for r in retrieved}
        cap = self.config.max_context_passages
        for r in retrieved:
            if len(by_id) >= cap:
                break
            for nb in self.store.fetch_neighbors(
                r.get("source", ""), r.get("seq", -1), self.config.neighbor_window
            ):
                if len(by_id) >= cap:
                    break
                by_id.setdefault(nb["id"], nb)
        # order by (source, seq) so passages read naturally
        return sorted(by_id.values(),
                      key=lambda x: (x.get("source", ""), x.get("seq", 0)))

    # ---------- Generation ----------

    def query(self, question: str, history: Optional[List[dict]] = None,
              return_sources: bool = True) -> Dict[str, Any]:
        search_q = self.rewrite_query(question, history=history)
        extra = [question] if self.config.enable_multi_query else None
        retrieved = self.retrieve(search_q, extra_queries=extra)
        contexts = self.format_contexts(self._expand_with_neighbors(retrieved))
        answer = self.llm.generate(
            query=question,                 # answer the ORIGINAL question
            contexts=contexts,
            temperature=self.config.temperature,
            max_tokens=self.config.max_tokens,
            history=history,
        )
        result = {"answer": answer}
        if return_sources:
            result["sources"] = retrieved   # primary hits only (clean sources list)
        return result
